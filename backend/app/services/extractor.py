import os
import re
from pathlib import Path
from typing import List, Optional, Tuple
import pypdf
import docx


class ExtractedPage:
    def __init__(self, page_number: int, raw_text: str, cleaned_text: str):
        self.page_number = page_number
        self.raw_text = raw_text
        self.cleaned_text = cleaned_text


class TextExtractorService:
    @staticmethod
    def clean_text(text: str) -> str:
        """Normalizes whitespace and removes unprintable artifacts."""
        if not text:
            return ""
        # Replace non-breaking spaces and special quotes
        cleaned = text.replace("\xa0", " ").replace("\r\n", "\n").replace("\r", "\n")
        # Replace curly quotes and apostrophes
        cleaned = cleaned.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
        # Normalize excessive newlines (keep at most 2)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        # Normalize spaces (keep line breaks)
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in cleaned.split("\n")]
        return "\n".join(lines).strip()

    @staticmethod
    def infer_year(filename: str, sample_text: Optional[str] = None) -> Optional[str]:
        """
        Attempts to detect the academic year or examination year from
        the filename or document header text.
        Looks for patterns like '2023', '2022/2023', '2021-2022', '2020_2021'.
        """
        # Search filename first
        combined = filename
        if sample_text:
            combined = f"{filename} {sample_text[:1000]}"

        # Pattern for 2022/2023 or 2022-2023 or 2022_2023
        session_match = re.search(r"\b(20[1-3][0-9])\s*[-/_]\s*(20[1-3][0-9]|\d{2})\b", combined)
        if session_match:
            return session_match.group(0)

        # Single year pattern 2010 to 2035
        year_match = re.search(r"\b(20[1-3][0-9])\b", combined)
        if year_match:
            return year_match.group(1)

        return None

    def extract_from_pdf(self, file_path: str) -> List[ExtractedPage]:
        """Extracts text page by page from a PDF file."""
        pages: List[ExtractedPage] = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for idx, page in enumerate(reader.pages):
                raw = page.extract_text() or ""
                cleaned = self.clean_text(raw)
                pages.append(ExtractedPage(page_number=idx + 1, raw_text=raw, cleaned_text=cleaned))
        return pages

    def extract_from_docx(self, file_path: str) -> List[ExtractedPage]:
        """Extracts text from a DOCX Word document."""
        doc = docx.Document(file_path)
        full_text_list = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text_list.append(paragraph.text)

        # Also extract table contents
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text_list.append(row_text)

        raw = "\n\n".join(full_text_list)
        cleaned = self.clean_text(raw)
        return [ExtractedPage(page_number=1, raw_text=raw, cleaned_text=cleaned)]

    def extract_from_text(self, file_path: str) -> List[ExtractedPage]:
        """Extracts text from a plain text or markdown file."""
        raw = ""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                raw = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                raw = f.read()

        cleaned = self.clean_text(raw)
        return [ExtractedPage(page_number=1, raw_text=raw, cleaned_text=cleaned)]

    def extract(self, file_path: str, filename: str) -> Tuple[List[ExtractedPage], Optional[str]]:
        """
        Main extraction entry point.
        Dispatches to the appropriate parser by file extension.
        Returns (pages, inferred_year).
        """
        ext = Path(filename).suffix.lower()
        pages: List[ExtractedPage] = []

        if ext == ".pdf":
            pages = self.extract_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            pages = self.extract_from_docx(file_path)
        elif ext in [".txt", ".md", ".markdown", ".csv"]:
            pages = self.extract_from_text(file_path)
        else:
            # Fallback attempt as text
            try:
                pages = self.extract_from_text(file_path)
            except Exception as e:
                raise ValueError(f"Unsupported file format '{ext}'. Supported formats: PDF, DOCX, TXT, MD.") from e

        # Infer year using filename and first page content
        sample_text = pages[0].cleaned_text if pages else ""
        inferred_year = self.infer_year(filename, sample_text)

        return pages, inferred_year


extractor_service = TextExtractorService()
